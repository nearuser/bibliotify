import spotipy
from spotipy.oauth2 import SpotifyOAuth
from http.server import BaseHTTPRequestHandler, HTTPServer
import webbrowser
from docx import Document

from dotenv import load_dotenv
import os

# Carga las variables de entorno desde el archivo .env
load_dotenv()

# Configura las credenciales de tu aplicación
client_id = os.getenv("CLIENT_ID")
client_secret = os.getenv("CLIENT_SECRET")
redirect_uri = os.getenv("REDIRECT_URI")

# Crea una instancia de SpotifyOAuth con tus credenciales y la URI de redireccionamiento
auth_manager = SpotifyOAuth(client_id=client_id,
                            client_secret=client_secret,
                            redirect_uri=redirect_uri,
                            scope="user-library-read",
                            open_browser=True)

# Abre la página de autorización de Spotify en el navegador predeterminado
auth_url = auth_manager.get_authorize_url()
print("Abriendo página de autorización de Spotify en el navegador predeterminado...")
webbrowser.open(auth_url)

# Clase para manejar la solicitud de callback
class CallbackHandler(BaseHTTPRequestHandler):
    def do_GET(self):
        self.send_response(200)
        self.send_header('Content-type', 'text/html')
        self.end_headers()

        # Maneja la respuesta de autorización de Spotify
        code = self.path.split('code=')[1] if 'code=' in self.path else None
        if code:
            # Intercambia el código de autorización por un token de acceso
            token_info = auth_manager.get_access_token(code)
            # El token de acceso está en token_info['access_token']

            # Crea una instancia de Spotipy utilizando el token de acceso
            sp = spotipy.Spotify(auth=token_info['access_token'])

            # Obtiene todas las pistas guardadas en la biblioteca del usuario
            library_results = sp.current_user_saved_tracks(limit=50)  # Puedes ajustar el límite según sea necesario
            tracks = library_results['items']

            # Itera sobre las siguientes páginas, si las hay
            while library_results['next']:
                library_results = sp.next(library_results)
                tracks.extend(library_results['items'])

            # Ordena las pistas por el nombre del artista
            tracks.sort(key=lambda x: x['track']['artists'][0]['name'])

            # Crea un nuevo archivo DOCX
            doc = Document()
            doc.add_heading('Listado de pistas ordenadas por artista', level=1)

            # Itera sobre las canciones y agrega cada una al documento
            for idx, item in enumerate(tracks):
                track = item['track']
                artist = track['artists'][0]['name']
                song = track['name']
                text = f"{artist} - {song}"
                doc.add_paragraph(text)

            # Guarda el archivo DOCX
            doc.save('listado_pistas.docx')

            response = "<html><body><h1>Autorizacion exitosa. Puedes cerrar esta ventana.</h1></body></html>"

            self.wfile.write(response.encode('utf-8'))  # Codifica la cadena a bytes utilizando UTF-8
        else:
            response = "<html><body><h1>Error de autorizacion. No se proporciono un codigo de autorizacion.</h1></body></html>"
            self.wfile.write(response.encode('utf-8'))  # Codifica la cadena a bytes utilizando UTF-8

# Inicializa el servidor HTTP para manejar la respuesta de callback
server_address = ('', 8888)
httpd = HTTPServer(server_address, CallbackHandler)

# Maneja las solicitudes de callback hasta que se interrumpa manualmente
print('Servidor HTTP iniciado. Esperando respuesta de autorización...')
httpd.handle_request()